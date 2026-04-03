"""Web interface for Carbon Footprint Optimization Engine (CfoE)."""

from __future__ import annotations

import asyncio
from queue import Queue
import csv
import json
import threading
from textwrap import wrap
from datetime import datetime, timezone
from functools import lru_cache
from pathlib import Path
from typing import Any
from uuid import uuid4

from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException, Request, WebSocket, WebSocketDisconnect
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field
from docx import Document
from reportlab.lib.pagesizes import A4

from agents.calculation_agent import calculate_carbon_score
from agents.policy_agent import enforce_policy_hitl
from orchestrators.root_coordinator import create_root_coordinator

from config.groq_config import get_groq_client

try:
    from groq import Groq
except ImportError:  # pragma: no cover
    Groq = None


BASE_DIR = Path(__file__).resolve().parent
WEB_DIR = BASE_DIR / "web"
STATIC_DIR = WEB_DIR / "static"
DATA_DIR = BASE_DIR / "data"
HISTORY_PATH = DATA_DIR / "audit_history.json"
PENDING_PATH = DATA_DIR / "pending_approvals.json"
OUTPUT_DIR = BASE_DIR / "outputs"
OUTPUT_CSV_PATH = OUTPUT_DIR / "audits_master.csv"

lock = threading.Lock()
log_queue = Queue()


class AuditRequest(BaseModel):
    supplier_name: str = Field(min_length=2, max_length=120)
    emissions: float = Field(ge=0)
    violations: int = Field(ge=0, le=500)
    notes: str = Field(default="", max_length=1000)


class AuditResponse(BaseModel):
    job_id: str
    audit_id: str
    timestamp: str
    supplier_name: str
    emissions: float
    violations: int
    notes: str
    risk_score: float
    classification: str
    emissions_score: float
    violations_score: float
    external_risk_score: float
    policy_decision: str
    human_approval_required: bool
    policy_reason: str
    recommended_action: str
    report_text: str
    report_source: str
    download_links: dict[str, str]
    status: str  # "completed" or "pending_approval"


class ApprovalRequest(BaseModel):
    audit_id: str
    decision: str  # "approve" or "reject"
    approver_name: str = Field(min_length=2, max_length=100)
    approval_notes: str = Field(default="", max_length=500)


app = FastAPI(
    title="CfoE Dashboard API",
    description="Interactive interface for supplier ESG risk audits.",
    version="1.0.0",
)


@app.middleware("http")
async def add_no_cache_headers(request: Request, call_next):
    response = await call_next(request)

    # Prevent stale frontend assets during local development.
    if request.url.path == "/" or request.url.path.startswith("/static/"):
        response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
        response.headers["Pragma"] = "no-cache"
        response.headers["Expires"] = "0"

    return response

# Ensure static output directory exists before mount.
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")
app.mount("/outputs", StaticFiles(directory=str(OUTPUT_DIR)), name="outputs")


@lru_cache(maxsize=1)
def get_client():
    load_dotenv()
    api_key = __import__("os").getenv("GROQ_API_KEY")
    if not api_key or Groq is None:
        return None
    try:
        return Groq(api_key=api_key)
    except Exception:
        return None


def ensure_storage() -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    if not HISTORY_PATH.exists():
        HISTORY_PATH.write_text("[]", encoding="utf-8")
    
    if not PENDING_PATH.exists():
        PENDING_PATH.write_text("[]", encoding="utf-8")

    if not OUTPUT_CSV_PATH.exists():
        with OUTPUT_CSV_PATH.open("w", encoding="utf-8", newline="") as csv_file:
            writer = csv.DictWriter(
                csv_file,
                fieldnames=[
                    "audit_id",
                    "timestamp",
                    "supplier_name",
                    "emissions",
                    "violations",
                    "risk_score",
                    "classification",
                    "policy_decision",
                    "human_approval_required",
                    "report_source",
                ],
            )
            writer.writeheader()


def load_history() -> list[dict[str, Any]]:
    ensure_storage()
    with lock:
        try:
            content = HISTORY_PATH.read_text(encoding="utf-8").strip() or "[]"
            data = json.loads(content)
            return data if isinstance(data, list) else []
        except json.JSONDecodeError:
            return []


def save_history(history: list[dict[str, Any]]) -> None:
    with lock:
        HISTORY_PATH.write_text(json.dumps(history, indent=2), encoding="utf-8")


def load_pending() -> list[dict[str, Any]]:
    ensure_storage()
    with lock:
        try:
            content = PENDING_PATH.read_text(encoding="utf-8").strip() or "[]"
            data = json.loads(content)
            return data if isinstance(data, list) else []
        except json.JSONDecodeError:
            return []


def save_pending(pending: list[dict[str, Any]]) -> None:
    with lock:
        PENDING_PATH.write_text(json.dumps(pending, indent=2), encoding="utf-8")


def make_audit_prompt(req: AuditRequest) -> str:
    return f"""
Conduct a comprehensive ESG audit for the following supplier:

Supplier Name: {req.supplier_name}
Annual CO2 Emissions: {req.emissions} tons
Regulatory Violations: {req.violations}
Additional Notes: {req.notes or 'N/A'}

Please provide a complete risk assessment and recommendations.
"""


def build_fallback_report(req: AuditRequest, risk_data: dict[str, Any], policy_data: dict[str, Any]) -> str:
    return (
        "Executive Summary\n"
        f"Supplier: {req.supplier_name}\n"
        f"Risk Score: {risk_data['risk_score']} ({risk_data['classification']})\n\n"
        "Key Findings\n"
        f"- Emissions contribution score: {risk_data['emissions_score']}\n"
        f"- Violations contribution score: {risk_data['violations_score']}\n"
        f"- Policy decision: {policy_data['decision']}\n\n"
        "Recommended Action\n"
        f"{policy_data['recommended_action']}\n"
    )


def run_audit(req: AuditRequest) -> dict[str, Any]:
    log_queue.put({"type": "info", "message": f"Starting audit for {req.supplier_name}..."})
    
    log_queue.put({"type": "info", "message": "[1/4] Calculating ESG risk scores..."})
    risk_data = calculate_carbon_score(emissions=req.emissions, violations=req.violations)
    log_queue.put({"type": "success", "message": f"✓ Risk Score: {risk_data['risk_score']} ({risk_data['classification']})"})
    
    log_queue.put({"type": "info", "message": "[2/4] Enforcing policy rules..."})
    policy_data = enforce_policy_hitl(risk_score=risk_data["risk_score"], supplier_name=req.supplier_name)
    log_queue.put({"type": "success", "message": f"✓ Policy Decision: {policy_data['decision']}"})

    report_source = "deterministic-fallback"
    report_text = build_fallback_report(req, risk_data, policy_data)
    external_risk_score = 0.0

    client = get_client()
    if client is not None:
        try:
            log_queue.put({"type": "info", "message": "[3/4] Generating AI report with multi-agent pipeline..."})
            coordinator = create_root_coordinator(client)
            response = coordinator.generate_content(make_audit_prompt(req))
            report_text = response.text
            
            # Extract external_risk_score from coordinator context if available
            if hasattr(coordinator, 'context') and 'external_risk_score' in coordinator.context.state:
                external_risk_score = coordinator.context.state.get('external_risk_score', 0.0)
            
            report_source = "groq-llama"
            log_queue.put({"type": "success", "message": "✓ AI report generated successfully"})
        except Exception as e:
            log_queue.put({"type": "warning", "message": f"⚠ AI report failed, using fallback: {str(e)[:100]}"})
            report_source = "deterministic-fallback"
    else:
        log_queue.put({"type": "warning", "message": "⚠ AI client unavailable, using deterministic report"})

    log_queue.put({"type": "info", "message": "[4/4] Finalizing audit results..."})
    result = {
        "job_id": f"JOB-{uuid4().hex[:8].upper()}",
        "audit_id": f"AUD-{uuid4().hex[:10].upper()}",
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "supplier_name": req.supplier_name,
        "emissions": req.emissions,
        "violations": req.violations,
        "notes": req.notes,
        "risk_score": risk_data["risk_score"],
        "classification": risk_data["classification"],
        "emissions_score": risk_data["emissions_score"],
        "violations_score": risk_data["violations_score"],
        "external_risk_score": external_risk_score,
        "policy_decision": policy_data["decision"],
        "human_approval_required": policy_data["human_approval_required"],
        "policy_reason": policy_data["reason"],
        "recommended_action": policy_data["recommended_action"],
        "report_text": report_text,
        "report_source": report_source,
        "download_links": {},
        "status": "pending_approval" if policy_data["human_approval_required"] else "completed",
    }
    
    # HITL Workflow Pause: If human approval required, save to pending queue
    if policy_data["human_approval_required"]:
        log_queue.put({"type": "warning", "message": "🚨 CRITICAL RISK - Audit paused for human approval"})
        result["status"] = "pending_approval"
        result["approval_status"] = "pending"
        result["approver_name"] = None
        result["approval_notes"] = None
        result["approval_timestamp"] = None
    else:
        log_queue.put({"type": "success", "message": f"✓ Audit complete for {req.supplier_name}"})
    
    return result


def _write_pdf(pdf_path: Path, result: dict[str, Any]) -> None:
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    
    doc = SimpleDocTemplate(str(pdf_path), pagesize=A4, topMargin=0.75*inch, bottomMargin=0.75*inch, leftMargin=0.75*inch, rightMargin=0.75*inch)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#0d7c66'),
        spaceAfter=12,
        alignment=TA_CENTER
    )
    story.append(Paragraph(f"CfoE Audit Report", title_style))
    story.append(Spacer(1, 0.2*inch))
    
    # Summary Table with word wrapping
    summary_data = [
        ['Audit ID:', Paragraph(result['audit_id'], styles['Normal'])],
        ['Job ID:', Paragraph(result['job_id'], styles['Normal'])],
        ['Timestamp:', Paragraph(result['timestamp'][:19], styles['Normal'])],
        ['Supplier:', Paragraph(result['supplier_name'], styles['Normal'])],
        ['Emissions:', Paragraph(f"{result['emissions']} tons CO2", styles['Normal'])],
        ['Violations:', Paragraph(str(result['violations']), styles['Normal'])],
        ['Risk Score:', Paragraph(f"{result['risk_score']} ({result['classification']})", styles['Normal'])],
    ]
    
    summary_table = Table(summary_data, colWidths=[1.5*inch, 4.5*inch])
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f0f0f0')),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
        ('ALIGN', (1, 0), (1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ('RIGHTPADDING', (0, 0), (-1, -1), 8),
    ]))
    story.append(summary_table)
    story.append(Spacer(1, 0.3*inch))
    
    # Policy Section
    policy_style = ParagraphStyle('PolicyHeading', parent=styles['Heading2'], fontSize=14, textColor=colors.HexColor('#0d7c66'))
    story.append(Paragraph("Policy Enforcement", policy_style))
    story.append(Spacer(1, 0.1*inch))
    
    policy_data = [
        ['Decision:', Paragraph(result['policy_decision'], styles['Normal'])],
        ['Reason:', Paragraph(result['policy_reason'], styles['Normal'])],
        ['Recommended Action:', Paragraph(result['recommended_action'], styles['Normal'])],
    ]
    
    policy_table = Table(policy_data, colWidths=[1.5*inch, 4.5*inch])
    policy_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f0f0f0')),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
        ('ALIGN', (1, 0), (1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ('RIGHTPADDING', (0, 0), (-1, -1), 8),
    ]))
    story.append(policy_table)
    story.append(Spacer(1, 0.3*inch))
    
    # Executive Report
    story.append(Paragraph("Executive Report", policy_style))
    story.append(Spacer(1, 0.1*inch))
    
    # Format report text with proper line breaks and structure
    report_style = ParagraphStyle(
        'ReportBody', 
        parent=styles['BodyText'], 
        fontSize=9, 
        leading=13,
        alignment=TA_LEFT,
        leftIndent=0,
        rightIndent=0,
        spaceAfter=6
    )
    
    # Split report into lines and format properly
    report_lines = result['report_text'].split('\n')
    for line in report_lines:
        line = line.strip()
        if line:
            # Replace special characters that might cause issues
            line = line.replace('━', '-')
            line = line.replace('•', '*')
            
            # Check if it's a section header (numbered or all caps)
            if line and (line[0].isdigit() or line.isupper() or line.startswith('---')):
                header_style = ParagraphStyle(
                    'SectionHeader',
                    parent=styles['Heading3'],
                    fontSize=10,
                    textColor=colors.HexColor('#0d7c66'),
                    spaceAfter=4,
                    spaceBefore=8,
                    fontName='Helvetica-Bold'
                )
                story.append(Paragraph(line, header_style))
            else:
                story.append(Paragraph(line, report_style))
    
    doc.build(story)


def _write_docx(docx_path: Path, result: dict[str, Any]) -> None:
    doc = Document()
    doc.add_heading(f"CfoE Audit Report - {result['audit_id']}", level=1)
    doc.add_paragraph(f"Job ID: {result['job_id']}")
    doc.add_paragraph(f"Timestamp: {result['timestamp']}")
    doc.add_paragraph(f"Supplier: {result['supplier_name']}")
    doc.add_paragraph(f"Emissions: {result['emissions']}")
    doc.add_paragraph(f"Violations: {result['violations']}")
    doc.add_paragraph(f"Risk Score: {result['risk_score']} ({result['classification']})")
    doc.add_paragraph(f"Policy Decision: {result['policy_decision']}")
    doc.add_paragraph(f"Policy Reason: {result['policy_reason']}")
    doc.add_paragraph(f"Recommended Action: {result['recommended_action']}")
    doc.add_heading("Executive Report", level=2)
    doc.add_paragraph(result["report_text"])
    doc.save(str(docx_path))


def export_audit_files(result: dict[str, Any]) -> dict[str, str]:
    ensure_storage()

    safe_stem = result["audit_id"].lower()
    job_dir = OUTPUT_DIR / result["job_id"].lower()
    job_dir.mkdir(parents=True, exist_ok=True)

    pdf_path = job_dir / f"{safe_stem}.pdf"
    docx_path = job_dir / f"{safe_stem}.docx"

    _write_pdf(pdf_path, result)
    _write_docx(docx_path, result)

    with lock:
        with OUTPUT_CSV_PATH.open("a", encoding="utf-8", newline="") as csv_file:
            writer = csv.DictWriter(
                csv_file,
                fieldnames=[
                    "audit_id",
                    "timestamp",
                    "supplier_name",
                    "emissions",
                    "violations",
                    "risk_score",
                    "classification",
                    "policy_decision",
                    "human_approval_required",
                    "report_source",
                ],
            )
            writer.writerow(
                {
                    "audit_id": result["audit_id"],
                    "timestamp": result["timestamp"],
                    "supplier_name": result["supplier_name"],
                    "emissions": result["emissions"],
                    "violations": result["violations"],
                    "risk_score": result["risk_score"],
                    "classification": result["classification"],
                    "policy_decision": result["policy_decision"],
                    "human_approval_required": result["human_approval_required"],
                    "report_source": result["report_source"],
                }
            )

    return {
        "pdf": f"/api/audits/{result['audit_id']}/pdf/download",
        "docx": f"/outputs/{result['job_id'].lower()}/{docx_path.name}",
    }


@app.on_event("startup")
def startup_event() -> None:
    ensure_storage()


@app.get("/")
def serve_index() -> FileResponse:
    return FileResponse(WEB_DIR / "index.html")


@app.post("/api/audit", response_model=AuditResponse)
def create_audit(payload: AuditRequest) -> dict[str, Any]:
    try:
        result = run_audit(payload)
        result["download_links"] = export_audit_files(result)
    except Exception as exc:  # pragma: no cover
        raise HTTPException(status_code=500, detail=f"Audit failed: {exc}") from exc

    # If human approval required, save to pending queue instead of history
    if result.get("human_approval_required", False):
        pending = load_pending()
        pending.insert(0, result)
        save_pending(pending)
    else:
        # Auto-approved audits go directly to history
        history = load_history()
        history.insert(0, result)
        save_history(history[:500])
    
    return result


@app.get("/api/audits")
def list_audits(limit: int = 100) -> dict[str, Any]:
    history = load_history()
    return {"items": history[:max(1, min(limit, 500))], "count": len(history)}


@app.get("/api/approvals")
def list_pending_approvals() -> dict[str, Any]:
    """Get all audits pending human approval"""
    pending = load_pending()
    return {"items": pending, "count": len(pending)}


@app.post("/api/approvals/{audit_id}/approve")
def approve_audit(audit_id: str, approval: ApprovalRequest) -> dict[str, Any]:
    """Approve a pending audit and move it to history"""
    if approval.audit_id != audit_id:
        raise HTTPException(status_code=400, detail="Audit ID mismatch")
    
    if approval.decision != "approve":
        raise HTTPException(status_code=400, detail="Use /reject endpoint for rejections")
    
    pending = load_pending()
    audit = next((x for x in pending if x.get("audit_id") == audit_id), None)
    
    if audit is None:
        raise HTTPException(status_code=404, detail="Pending audit not found")
    
    # Update audit with approval info
    audit["status"] = "completed"
    audit["approval_status"] = "approved"
    audit["approver_name"] = approval.approver_name
    audit["approval_notes"] = approval.approval_notes
    audit["approval_timestamp"] = datetime.now(timezone.utc).isoformat()
    
    # Remove from pending and add to history
    pending = [x for x in pending if x.get("audit_id") != audit_id]
    save_pending(pending)
    
    history = load_history()
    history.insert(0, audit)
    save_history(history[:500])
    
    return {"status": "approved", "audit": audit}


@app.post("/api/approvals/{audit_id}/reject")
def reject_audit(audit_id: str, approval: ApprovalRequest) -> dict[str, Any]:
    """Reject a pending audit"""
    if approval.audit_id != audit_id:
        raise HTTPException(status_code=400, detail="Audit ID mismatch")
    
    if approval.decision != "reject":
        raise HTTPException(status_code=400, detail="Use /approve endpoint for approvals")
    
    pending = load_pending()
    audit = next((x for x in pending if x.get("audit_id") == audit_id), None)
    
    if audit is None:
        raise HTTPException(status_code=404, detail="Pending audit not found")
    
    # Update audit with rejection info
    audit["status"] = "rejected"
    audit["approval_status"] = "rejected"
    audit["approver_name"] = approval.approver_name
    audit["approval_notes"] = approval.approval_notes
    audit["approval_timestamp"] = datetime.now(timezone.utc).isoformat()
    
    # Remove from pending and add to history (with rejected status)
    pending = [x for x in pending if x.get("audit_id") != audit_id]
    save_pending(pending)
    
    history = load_history()
    history.insert(0, audit)
    save_history(history[:500])
    
    return {"status": "rejected", "audit": audit}


@app.delete("/api/audits")
def clear_audits() -> dict[str, Any]:
    save_history([])
    return {"status": "ok"}


@app.delete("/api/approvals")
def clear_pending_approvals() -> dict[str, Any]:
    """Clear all pending approvals"""
    save_pending([])
    return {"status": "ok"}


@app.get("/api/metrics")
def metrics() -> dict[str, Any]:
    history = load_history()
    if not history:
        return {
            "total_audits": 0,
            "avg_risk_score": 0,
            "critical_rate": 0,
            "classifications": {"Low Risk": 0, "Moderate Risk": 0, "Critical Risk": 0},
        }

    total = len(history)
    avg_score = sum(item["risk_score"] for item in history) / total
    critical = sum(1 for item in history if item["classification"] == "Critical Risk")
    counts = {"Low Risk": 0, "Moderate Risk": 0, "Critical Risk": 0}
    for item in history:
        counts[item["classification"]] = counts.get(item["classification"], 0) + 1

    return {
        "total_audits": total,
        "avg_risk_score": round(avg_score, 3),
        "critical_rate": round((critical / total) * 100, 1),
        "classifications": counts,
    }


@app.get("/api/audits/{audit_id}/pdf/view")
def view_pdf(audit_id: str) -> FileResponse:
    history = load_history()
    item = next((x for x in history if x.get("audit_id") == audit_id), None)
    if item is None:
        raise HTTPException(status_code=404, detail="Audit not found")

    job_id = item.get("job_id", "")
    if not job_id:
        raise HTTPException(status_code=404, detail="Job export not found")

    pdf_name = f"{audit_id.lower()}.pdf"
    pdf_path = OUTPUT_DIR / job_id.lower() / pdf_name
    if not pdf_path.exists():
        raise HTTPException(status_code=404, detail="PDF file not found")

    return FileResponse(
        str(pdf_path),
        media_type="application/pdf",
        headers={"Content-Disposition": f"inline; filename={pdf_name}"},
    )


@app.get("/api/audits/{audit_id}/pdf/download")
def download_pdf(audit_id: str) -> FileResponse:
    history = load_history()
    item = next((x for x in history if x.get("audit_id") == audit_id), None)
    if item is None:
        raise HTTPException(status_code=404, detail="Audit not found")

    job_id = item.get("job_id", "")
    if not job_id:
        raise HTTPException(status_code=404, detail="Job export not found")

    pdf_name = f"{audit_id.lower()}.pdf"
    pdf_path = OUTPUT_DIR / job_id.lower() / pdf_name
    if not pdf_path.exists():
        raise HTTPException(status_code=404, detail="PDF file not found")

    return FileResponse(
        str(pdf_path),
        media_type="application/pdf",
        filename=pdf_name,
        headers={"Content-Disposition": f"attachment; filename={pdf_name}"},
    )


@app.websocket("/ws/logs")
async def websocket_logs(websocket: WebSocket):
    await websocket.accept()
    try:
        while True:
            # Check for new log messages
            if not log_queue.empty():
                log_msg = log_queue.get()
                await websocket.send_json(log_msg)
            await asyncio.sleep(0.1)
    except WebSocketDisconnect:
        pass
